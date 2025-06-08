# src/report_generator.py
from docx import Document
from matplotlib import pyplot as plt
import pandas as pd

def export_summary_to_word(text, path="stock_summary.docx"):
    doc = Document()
    doc.add_heading("📊 Stock Summary Report", level=1)
    doc.add_paragraph(text)
    doc.save(path)
    return path

def export_plot_to_pdf(df, path="plot.pdf"):
    df.plot(y=["close"])
    plt.title("Close Price History")
    plt.xlabel("Time")
    plt.ylabel("Price")
    plt.tight_layout()
    plt.savefig(path)
    return path
